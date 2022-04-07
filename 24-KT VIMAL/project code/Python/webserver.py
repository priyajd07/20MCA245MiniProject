import os

import datetime

import cv2
import pytesseract
from PIL import Image
from flask import *

from werkzeug.utils import secure_filename
import speech_recognition as sr

from src.dbconnect import *
app=Flask(__name__)
import docx
from docx.shared import Pt


def docc1(text,ftype,fface,fsize):
    doc = docx.Document()

    text=text.split('\\n')
    print(len(text))
    for i in text:
        doc_para = doc.add_paragraph('')

        if ftype=='bold':
            doc_para.add_run(i).bold = True
        elif ftype=='italic':
            doc_para.add_run(i).italic = True
        elif ftype == 'all':
            d=doc_para.add_run(i)
            d.bold = True
            d.italic = True
            print("okkkkkkkkkkkkkk")

        else:
            doc_para.add_run(i)

    Tstyle = doc.styles['Normal']
    font = Tstyle.font
    font.name = fface
    font.size = Pt(fsize)
    import time

    fn = time.strftime("%Y%m%d_%H%M%S") + "." + ".doc"


    doc.save('static/texttodoc/'+fn)
    return fn
@app.route('/imgtotext2',methods=['post'])
def imgtotext2():
    print(request.form)
    text=request.form['text']
    ftype=request.form['ftype']
    fface=request.form['fface']
    fsize=request.form['fsize']
    id=request.form['id']


    res=docc1(text,ftype,fface,int(fsize))
    print("rrrrrr",res)
    qry="insert into history values(null,%s,%s,curdate())"
    val=(id,res)
    iud(qry,val)
    return jsonify({"task":res})
@app.route('/imgtotext1',methods=['post'])
def imgtotext1():
    text=request.form['text']
    ftype=request.form['ftype']
    fface=request.form['fface']
    fsize=request.form['fsize']
    image = request.files['files']
    image.save("static/texttodoc/sampe.png")
    print(image)
    img=cv2.imread("static/texttodoc/sampe.png")
    imgg= cv2.resize(img, (400, 400),
               interpolation = cv2.INTER_NEAREST)
    cv2.imwrite("static/texttodoc/sampe1.png",imgg)
    res=docc(text,ftype,fface,int(fsize))

    return jsonify({"task":res})
def docc(text,ftype,fface,fsize):
    doc = docx.Document()

    text=text.split('\\n')
    print(len(text))
    for i in text:
        doc_para = doc.add_paragraph('')

        if ftype=='bold':
            doc_para.add_run(i).bold = True
        elif ftype=='italic':
            doc_para.add_run(i).italic = True
        elif ftype == 'all':
            d=doc_para.add_run(i)
            d.bold = True
            d.italic = True
            print("okkkkkkkkkkkkkk")

        else:
            doc_para.add_run(i)
    doc.add_picture("static/texttodoc/sampe1.png")
    Tstyle = doc.styles['Normal']
    font = Tstyle.font
    font.name = fface
    font.size = Pt(fsize)
    fn = datetime.now().strftime("%Y%m%d%H%M%S") + ".doc"

    doc.save('static/texttodoc/'+fn)
    return fn
@app.route("/logincode",methods=['post'])
def logincode():
    uname=request.form['username']
    password=request.form['password']
    q="select * from login where username=%s and password=%s "
    val=uname,password
    res=selectonecond(q,val)
    if res is None:
        return jsonify({'task': 'invalid'})
    else:
        return jsonify({'task': 'success','lid':res[0],'type':res[3]})




@app.route("/register",methods=['post'])
def register():
    fname=request.form['fname']
    lname=request.form['lname']
    Phone=request.form['Phone']
    Place=request.form['Place']
    Email=request.form['Email']
    username=request.form['username']
    password=request.form['password']
    qry = "insert into login values(null,%s,%s,'user')"
    value = (username, password)
    lid = iud(qry, value)
    qry1 = "insert into user values(null,%s,%s,%s,%s,%s,%s)"
    val = (str(lid), fname, lname,Place,Phone,Email)
    iud(qry1, val)

    return jsonify({'task': 'success'})

@app.route('/doctospeech',methods=['post'])
def doctospeech():


    file = request.files['file']
    ff = secure_filename(file.filename)


    import time

    req = time.strftime("%Y%m%d_%H%M%S") + "." +".doc"
    print(req)
    file.save(os.path.join('./static/doctospeech', req))

    import docx2txt
    result = docx2txt.process("./static/doctospeech/"+req)
    print(result)
    return jsonify({'task':result})


@app.route('/imgtotext',methods=['post'])
def imgtotext():
    print(request.form)
    image = request.files['files']
    print(image)
    img = secure_filename(image.filename)

    image.save(os.path.join("./static/imgtotext", img))
    text=main(os.path.join("./static/imgtotext", img))
    print(text,"===========================================")
    return jsonify(text)
def main(path):
    # Get File Name from Command Line
    # path = input("Enter the file path : ").strip()

    # Load the required image
    print("path===",path)
    image = cv2.imread(path)
    # print(image)
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)

    # Store grayscale image as a temporary file to apply OCR
    filename = "{}.png".format("temp")
    cv2.imwrite(filename, gray)

    # Load the image as a PIL/Pillow image, apply OCR, and then delete the temporary file
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    text = pytesseract.image_to_string(Image.open(filename))

    print("OCR Text is " + text.strip())
    # val=text.strip().split('D')
    # atte=val[0].split(',')
    # print(atte[0],atte[1])
    # # print(val[1])
    return  text.strip()


@app.route('/photoupload', methods=['POST','GET'])
def photoupload():
    try:
        # print(request.form)
        print(request.files)
        # lid = request.form['lid']
        # print(lid)

        file=request.files['file']
        ff=secure_filename(file.filename)
        print(ff)
        fl = ff.split('.')
        print(fl[1])
        import time
        ffl=time.strftime("%Y%m%d_%H%M%S")
        req = time.strftime("%Y%m%d_%H%M%S") + "." + fl[1]
        file.save(os.path.join('./static/audio', req))
        #
        #
        # fl=fl.split('.')
        # print(fl[len(fl)-1])
        # print(datetime.now().strftime("%Y%m%d%H%M%S"))
        # fl=datetime.now().strftime("%Y%m%d%H%M%S")+"."+fl[len(fl)-1]
        # print("llllllllllllllllllllll",fl)
        # file.save(os.path.join("static/audio/",fl))
        # os.system('ffmpeg -i C:\\RINJU\\textaudio\src\static\\audio\\'+fl+' F:\\RINJU\\textaudio\\src\\static\\audio\\'+fl+".wav")

        os.system('ffmpeg -i E:\\final\\OCR\\src\\static\\audio\\'+req+' E:\\final\\OCR\\src\\static\\audio\\'+ffl+".wav")
        #
        res=silence_based_conversion(ffl)
        print(res)

        # os.system('ffmpeg -i path/to/3gp.3gp path/to/wav.wav')
        # q = "insert into audio values(null,%s,%s)"
        # val = (str(lid), fl)
        # iud(q, val)
        # print("result==="+res)
        return jsonify({'task': res})
        # return "ok"


    except Exception as e:
        print(e)

        return jsonify({'task':"No result"})


def silence_based_conversion(fl):
    path = r'static/audio/'+fl+'.wav'



    r = sr.Recognizer()
    file=path
    # recognize the chunk
    with sr.AudioFile(file) as source:
        # remove this if it is not working
        # correctly.
        r.adjust_for_ambient_noise(source)
        audio_listened = r.listen(source)

    try:
        # try converting it to text
        rec = r.recognize_google(audio_listened)
        # write the output to the file.
        print(rec)
        return rec


    except sr.UnknownValueError:
        print("Could not understand audio")

    except sr.RequestError as e:
        print("Could not request results. check your internet connection")

@app.route('/texttodoc',methods=['post'])
def texttodoc():
    text= request.form['text']
    print(text)
    import docx
    doc = docx.Document()
    doc_para = doc.add_paragraph(text)
    print(doc_para)

    import time

    req = time.strftime("%Y%m%d_%H%M%S") + "." +".doc"
    doc.save(os.path.join('./static/texttodoc/', req))
    print(req)

    # fn = datetime.now().strftime("%Y%m%d%H%M%S") + ".doc"
    # doc.save(os.path.join("./static/texttodoc/", fn))
    # result=fn
    # pip install python-docx
    return jsonify({'task': req})




@app.route("/history",methods=['post'])
def history():
    print(request.form)
    uid=request.form['uid']
    qry="SELECT * from `history`  where uid=%s"
    val=uid
    res = androidselectall(qry,val)
    print(res)
    return jsonify(res)

if __name__=='__main__':
    app.run(host='0.0.0.0',port=5000)